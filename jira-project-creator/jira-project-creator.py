#!/usr/bin/env python3
"""
Jira Project Creator from Specification
Automatically creates Jira Epics, dependencies, timeline, and Plan from a spec document
"""

import argparse
import json
import sys
import subprocess
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Optional
import re
import os

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import anthropic
except ImportError:
    print("ERROR: anthropic not installed. Run: pip install anthropic")
    sys.exit(1)


class SpecParser:
    """Parse project specification document"""

    def __init__(self, spec_path: str):
        self.spec_path = spec_path
        self.doc = Document(spec_path)

    def extract_project_info(self) -> Dict:
        """Extract project name, timeline, scope"""
        info = {
            "name": "",
            "timeline_weeks": 0,
            "stages": []
        }

        full_text = "\n".join([p.text for p in self.doc.paragraphs])

        # Extract project name (look for "Project:" or similar)
        name_match = re.search(r'Project[:\s]+([^\n]+)', full_text, re.IGNORECASE)
        if name_match:
            info["name"] = name_match.group(1).strip()

        # Extract timeline (look for "X weeks" or "X-Y weeks")
        timeline_match = re.search(r'(\d+)[-–]?(\d+)?\s*weeks?', full_text, re.IGNORECASE)
        if timeline_match:
            info["timeline_weeks"] = int(timeline_match.group(2) or timeline_match.group(1))

        return info

    def extract_stages(self) -> List[Dict]:
        """Extract stages/epics from specification"""
        stages = []
        current_stage = None

        for para in self.doc.paragraphs:
            text = para.text.strip()

            # Look for stage headings (e.g., "Stage 1:", "EPIC-001:", etc.)
            stage_match = re.match(r'(?:Stage|Epic|STAGE|EPIC)[\s-]*(\d+)[:\s]+(.+)', text, re.IGNORECASE)
            if stage_match:
                if current_stage:
                    stages.append(current_stage)

                current_stage = {
                    "number": int(stage_match.group(1)),
                    "name": stage_match.group(2).strip(),
                    "description": "",
                    "tasks": [],
                    "timeline": "",
                    "dependencies": []
                }
            elif current_stage:
                # Add to description or tasks
                if text.startswith("- [ ]") or text.startswith("-"):
                    task = text.lstrip("- [ ]").strip()
                    current_stage["tasks"].append(task)
                elif "week" in text.lower() or "day" in text.lower():
                    current_stage["timeline"] = text
                else:
                    current_stage["description"] += text + "\n"

        if current_stage:
            stages.append(current_stage)

        return stages


class JiraClient:
    """Interact with Jira API"""

    def __init__(self, config: Dict):
        self.base_url = config["jira"]["instanceUrl"]
        self.auth = f"{config['jira']['email']}:{config['jira']['apiToken']}"
        self.project_key = config["jira"]["projectKey"]

    def create_epic(self, summary: str, description: str = "", labels: List[str] = None,
                   priority: str = "1", start_date: str = None, due_date: str = None) -> str:
        """Create an Epic in Jira"""

        fields = {
            "project": {"key": self.project_key},
            "summary": summary,
            "issuetype": {"id": "10000"},  # Epic
            "priority": {"id": priority}
        }

        if description:
            fields["description"] = {
                "type": "doc",
                "version": 1,
                "content": [{
                    "type": "paragraph",
                    "content": [{"type": "text", "text": description}]
                }]
            }

        if labels:
            fields["labels"] = labels

        if due_date:
            fields["duedate"] = due_date

        data = json.dumps({"fields": fields})

        result = subprocess.run(
            ["curl", "-s", "-u", self.auth, "-X", "POST",
             "-H", "Content-Type: application/json",
             "-d", data,
             f"{self.base_url}/rest/api/3/issue"],
            capture_output=True,
            text=True
        )

        response = json.loads(result.stdout)
        return response.get("key", "")

    def add_start_date(self, issue_key: str, start_date: str):
        """Add start date to an issue"""
        data = json.dumps({"fields": {"customfield_10015": start_date}})

        subprocess.run(
            ["curl", "-s", "-u", self.auth, "-X", "PUT",
             "-H", "Content-Type: application/json",
             "-d", data,
             f"{self.base_url}/rest/api/3/issue/{issue_key}"],
            capture_output=True
        )

    def create_dependency(self, blocker_key: str, blocked_key: str):
        """Create a 'blocks' dependency between two issues"""
        data = json.dumps({
            "type": {"name": "Blocks"},
            "inwardIssue": {"key": blocker_key},
            "outwardIssue": {"key": blocked_key}
        })

        subprocess.run(
            ["curl", "-s", "-u", self.auth, "-X", "POST",
             "-H", "Content-Type: application/json",
             "-d", data,
             f"{self.base_url}/rest/api/3/issueLink"],
            capture_output=True
        )


class ImplementationPlanGenerator:
    """Generate implementation plan from specification using Claude API"""

    def __init__(self, spec_path: str, api_key: str = None):
        self.spec_path = spec_path
        self.api_key = api_key or os.environ.get("ANTHROPIC_API_KEY")
        if not self.api_key:
            raise ValueError("ANTHROPIC_API_KEY not found in environment or config")
        self.client = anthropic.Anthropic(api_key=self.api_key)

    def extract_text_from_docx(self) -> str:
        """Extract full text from Word document"""
        doc = Document(self.spec_path)
        return "\n".join([p.text for p in doc.paragraphs])

    def generate_plan(self) -> List[Dict]:
        """Use Claude to generate implementation plan with stages"""
        spec_text = self.extract_text_from_docx()

        prompt = f"""Analyze this software specification and create a detailed implementation plan.

Specification:
{spec_text[:50000]}  # Limit to avoid token issues

CRITICAL REQUIREMENTS - You MUST explicitly identify and create separate stages for:

1. **Safety & Compliance Features**
   - Self-harm detection systems
   - Crisis intervention workflows
   - Mental health safety protocols
   - Medical/health compliance (HIPAA, etc.)
   - Content moderation
   - User safety mechanisms

2. **AI/ML Model Development**
   - Voice analysis models
   - Sentiment analysis models
   - Emotion detection systems
   - Open-source ML models (if mentioned)
   - Model training pipelines
   - AI recommendation engines
   - Any machine learning infrastructure

3. **Core Infrastructure**
   - Foundation/infrastructure setup
   - Authentication & authorization
   - Database architecture
   - API development

4. **User-Facing Features**
   - Onboarding flows
   - User profiles
   - Assessment systems
   - Content delivery

5. **Quality & Launch**
   - Testing & QA
   - Security hardening
   - Performance optimization
   - Launch preparation

Create a structured implementation plan with stages/epics. For each stage, provide:
- Stage number (1-N)
- Stage name (descriptive, specific title)
- Brief description (what gets built and why)
- Estimated timeline (in weeks)
- List of specific tasks
- Dependencies on other stages

Format your response as JSON array:
[
  {{
    "number": 1,
    "name": "Foundation & Authentication",
    "description": "Set up infrastructure and user authentication",
    "timeline": "Week 1-2",
    "tasks": ["AWS setup", "PostgreSQL installation", "Cognito configuration"],
    "dependencies": []
  }},
  ...
]

IMPORTANT GUIDELINES:
- Do NOT combine safety features with other stages - they deserve dedicated focus
- Do NOT hide ML model development inside generic stages - make them explicit
- Look for keywords: "detection", "analysis", "model", "ML", "AI", "safety", "crisis", "harm"
- If the spec mentions wellness/mental health, include safety monitoring stages
- Separate open-source components from proprietary ones
- **Create exactly 12-15 stages** - combine related work into cohesive stages
- Group frontend work into 1-2 stages maximum (not 6 separate stages)
- Dependencies should be SIMPLE and SEQUENTIAL:
  * Foundation stage (1) should have NO dependencies (it comes first)
  * Most stages depend only on the previous stage
  * Foundation stage blocks everything (include it in all other stages' dependencies)
  * Testing/QA stage comes last and depends on all feature stages
- Do NOT create circular dependencies (if A blocks B, then B cannot block A)
"""

        response = self.client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=16000,
            messages=[{"role": "user", "content": prompt}]
        )

        # Extract JSON from response
        response_text = response.content[0].text

        # Find JSON array in response
        json_match = re.search(r'\[[\s\S]*\]', response_text)
        if json_match:
            stages = json.loads(json_match.group(0))
            return stages
        else:
            raise ValueError("Could not extract JSON from Claude response")


class ProjectCreator:
    """Main orchestrator for creating Jira project from spec"""

    def __init__(self, spec_path: str, config: Dict, options: Dict = None):
        self.spec_path = spec_path
        self.jira = JiraClient(config)
        self.config = config
        self.options = options or {}
        self.plan_generator = None

        # Initialize plan generator if API key available
        if config.get("anthropic", {}).get("apiKey"):
            self.plan_generator = ImplementationPlanGenerator(
                spec_path,
                config["anthropic"]["apiKey"]
            )

    def calculate_dates(self, stage_number: int, total_stages: int) -> tuple:
        """Calculate start and due dates for a stage"""
        start_date_str = self.config["timeline"]["startDate"]
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d")

        compression = self.config["timeline"].get("compressionFactor", 1)

        # Simple calculation: each stage gets equal time
        days_per_stage = max(1, 7 // compression)

        stage_start = start_date + timedelta(days=(stage_number - 1) * days_per_stage)
        stage_end = stage_start + timedelta(days=days_per_stage - 1)

        return stage_start.strftime("%Y-%m-%d"), stage_end.strftime("%Y-%m-%d")

    def infer_dependencies(self, stages: List[Dict]) -> Dict[int, List[int]]:
        """Infer dependencies between stages

        Returns: Dict mapping stage_number -> list of stage numbers that block it
        """
        dependencies = {}

        # Stage 1 (Foundation) has no dependencies
        if len(stages) > 0:
            dependencies[stages[0]["number"]] = []

        # All other stages depend on:
        # 1. The previous stage (sequential)
        # 2. Stage 1 (foundation) if not already included
        for i, stage in enumerate(stages):
            if i > 0:
                stage_num = stage["number"]
                prev_stage_num = stages[i-1]["number"]

                # Initialize with previous stage
                dependencies[stage_num] = [prev_stage_num]

                # Also depend on foundation (stage 1) if not the immediate previous
                foundation_num = stages[0]["number"]
                if prev_stage_num != foundation_num:
                    dependencies[stage_num].append(foundation_num)

        return dependencies

    def create_project(self, dry_run: bool = False):
        """Create the complete Jira project"""

        print("=== Generating implementation plan from specification...")

        if self.plan_generator:
            # Use AI to generate plan
            stages = self.plan_generator.generate_plan()
            print(f"   AI generated {len(stages)} stages")
        else:
            # Fallback to parsing existing plan
            print("   WARNING: No API key, falling back to spec parsing")
            project_info = SpecParser(self.spec_path).extract_project_info()
            stages = SpecParser(self.spec_path).extract_stages()
            print(f"   Project: {project_info.get('name', 'Unknown')}")

        print(f"   Total Stages: {len(stages)}")
        print()

        if dry_run:
            print("=== DRY RUN MODE - No changes will be made")
            print()

        created_epics = {}

        print("=== Creating Epics...")
        for stage in stages:
            summary = f"STAGE-{stage['number']:03d}: {stage['name']}"

            # Create labels
            labels = [
                f"stage-{stage['number']:03d}",
                "stage"
            ]

            # Calculate dates
            start_date, due_date = self.calculate_dates(stage["number"], len(stages))

            if dry_run:
                print(f"   [DRY RUN] Would create: {summary}")
                print(f"             Dates: {start_date} to {due_date}")
                created_epics[stage["number"]] = f"AURA-{stage['number']}"
            else:
                epic_key = self.jira.create_epic(
                    summary=summary,
                    description=stage["description"][:500],
                    labels=labels,
                    due_date=due_date
                )

                if epic_key:
                    self.jira.add_start_date(epic_key, start_date)
                    created_epics[stage["number"]] = epic_key
                    print(f"   ✓ Created {epic_key}: {summary}")
                else:
                    print(f"   ✗ Failed to create: {summary}")

        print()

        if self.options.get("addDependencies", True):
            print("=== Creating Dependencies...")
            dependencies = self.infer_dependencies(stages)

            for blocked_num, blockers in dependencies.items():
                blocked_key = created_epics.get(blocked_num)
                for blocker_num in blockers:
                    blocker_key = created_epics.get(blocker_num)
                    if blocker_key and blocked_key:
                        if dry_run:
                            print(f"   [DRY RUN] {blocker_key} blocks {blocked_key}")
                        else:
                            self.jira.create_dependency(blocker_key, blocked_key)
                            print(f"   ✓ {blocker_key} blocks {blocked_key}")

        print()
        print("DONE Project creation complete!")
        print()
        print(f"=== View your project: {self.jira.base_url}/jira/software/projects/{self.jira.project_key}")
        print(f"=== View roadmap: {self.jira.base_url}/jira/software/c/projects/{self.jira.project_key}/roadmap")


def main():
    parser = argparse.ArgumentParser(description="Create Jira project from specification")
    parser.add_argument("--spec", required=True, help="Path to specification document (.docx)")
    parser.add_argument("--config", required=True, help="Path to config JSON file")
    parser.add_argument("--compress-timeline", type=int, help="Timeline compression factor (e.g., 7 for 1 week = 1 day)")
    parser.add_argument("--start-date", help="Project start date (YYYY-MM-DD or 'tomorrow')")
    parser.add_argument("--dry-run", action="store_true", help="Preview without creating")
    parser.add_argument("--output", help="Export summary to JSON file")

    args = parser.parse_args()

    # Load config
    with open(args.config) as f:
        config_text = f.read()

    # Substitute environment variables
    config_text = config_text.replace("${ANTHROPIC_API_KEY}", os.environ.get("ANTHROPIC_API_KEY", ""))
    config = json.loads(config_text)

    # Override config with command-line args
    if args.compress_timeline:
        config["timeline"]["compressionFactor"] = args.compress_timeline

    if args.start_date:
        if args.start_date.lower() == "tomorrow":
            config["timeline"]["startDate"] = (datetime.now() + timedelta(days=1)).strftime("%Y-%m-%d")
        else:
            config["timeline"]["startDate"] = args.start_date

    # Create project
    creator = ProjectCreator(args.spec, config, config.get("options", {}))
    creator.create_project(dry_run=args.dry_run)


if __name__ == "__main__":
    main()
